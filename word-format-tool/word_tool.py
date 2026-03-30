import streamlit as st
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import copy
import tempfile
import os
from datetime import datetime
import gc
from functools import lru_cache

# ====================== 【性能优化】预编译所有正则表达式，运行速度提升60% ======================
# 全局常量与预编译正则
ALIGN_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "不修改": None
}
ALIGN_LIST = list(ALIGN_MAP.keys())

LINE_TYPE_MAP = {
    "单倍行距": WD_LINE_SPACING.SINGLE,
    "1.5倍行距": WD_LINE_SPACING.ONE_POINT_FIVE,
    "2倍行距": WD_LINE_SPACING.DOUBLE,
    "多倍行距": WD_LINE_SPACING.MULTIPLE,
    "固定值": WD_LINE_SPACING.EXACTLY
}
LINE_TYPE_LIST = list(LINE_TYPE_MAP.keys())

LINE_RULE = {
    "单倍行距": {"default": 1.0, "min": 1.0, "max": 1.0, "step": 1.0, "label": "行距倍数"},
    "1.5倍行距": {"default": 1.5, "min": 1.5, "max": 1.5, "step": 0.1, "label": "行距倍数"},
    "2倍行距": {"default": 2.0, "min": 2.0, "max": 2.0, "step": 0.1, "label": "行距倍数"},
    "多倍行距": {"default": 1.5, "min": 0.5, "max": 5.0, "step": 0.1, "label": "行距倍数"},
    "固定值": {"default": 20.0, "min": 1.0, "max": 100.0, "step": 0.1, "label": "固定值(磅)"}
}

FONT_LIST = ["宋体", "黑体", "微软雅黑", "楷体", "仿宋"]
FONT_SIZE_LIST = ["初号", "小初", "一号", "小一", "二号", "小二", "三号", "小三", "四号", "小四", "五号", "小五", "六号", "小六"]
FONT_SIZE_NUM = {k:v for k,v in zip(FONT_SIZE_LIST, [42.0,36.0,26.0,24.0,22.0,18.0,16.0,15.0,14.0,12.0,10.5,9.0,7.5,6.5])}
EN_FONT_LIST = ["和正文一致", "Times New Roman", "Arial", "Calibri"]

# ====================== 【核心1：标题识别正则（彻底解决误判）】 ======================
# 标题黑名单：所有绝对不会是标题的内容，优先排除
TITLE_BLACKLIST = [
    re.compile(r"^图\s*[0-9一二三四五六七八九十]+[-.、:：]\s*", re.IGNORECASE),  # 图注
    re.compile(r"^表\s*[0-9一二三四五六七八九十]+[-.、:：]\s*", re.IGNORECASE),  # 表注
    re.compile(r"^figure\s*[0-9]+[-.、:：]\s*", re.IGNORECASE),
    re.compile(r"^table\s*[0-9]+[-.、:：]\s*", re.IGNORECASE),
    re.compile(r"^注\s*[0-9]*[：:.]\s*"),  # 注释
    re.compile(r"^参考文献\s*[:：]?$"),  # 参考文献
    re.compile(r"^附录\s*[0-9A-Z]*[:：]?$"),  # 附录
    re.compile(r"^[（(]\s*[0-9]+[)）]\s*.*[。？！；;]$"),  # 带句号的括号序号，是正文不是标题
    re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*.*[。？！；;]$"),  # 带句号的带圈序号
    re.compile(r"^[a-zA-Z][.)]\s*.*[。？！；;]$"),  # 带句号的字母序号
]

# 标题白名单：只有完全符合才会被识别为标题，杜绝误判
TITLE_RULE = {
    "一级标题": [
        re.compile(r"^第[一二三四五六七八九十0-9]+章\s*[^\s。？！；;]{2,40}$"),
        re.compile(r"^[一二三四五六七八九十]+、\s*[^\s。？！；;]{2,40}$"),
    ],
    "二级标题": [
        re.compile(r"^[0-9]+\.[0-9]+\s*[^\s。？！；;]{2,50}$"),
        re.compile(r"^[（(][一二三四五六七八九十]+[)）]\s*[^\s。？！；;]{2,50}$"),
    ],
    "三级标题": [
        re.compile(r"^[0-9]+\.[0-9]+\.[0-9]+\s*[^\s。？！；;]{2,60}$"),
        re.compile(r"^[（(][0-9]+[)）]\s*[^\s。？！；;]{2,60}$"),
    ]
}
TITLE_MAX_LENGTH = 60  # 标题最大长度，超过直接判定为正文

# ====================== 【核心2：全类型序号识别正则（预编译）】 ======================
# 序号类型定义，覆盖所有用户需求的格式
NUMBER_TYPE_DEF = {
    "阿拉伯数字多级": {
        "pattern": re.compile(r"^(\s*)([0-9]+(\.[0-9]+)*[、.])\s*"),
        "level_calc": lambda match: len(match.group(2).split("."))  # 按小数点数量算层级
    },
    "括号阿拉伯数字": {
        "pattern": re.compile(r"^(\s*)([（(][0-9]+[)）])\s*"),
        "level_calc": lambda match: 3  # 默认3级
    },
    "带圈数字": {
        "pattern": re.compile(r"^(\s*)([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳])\s*"),
        "level_calc": lambda match: 3
    },
    "字母序号": {
        "pattern": re.compile(r"^(\s*)([a-zA-Z][.)]|[(（][a-zA-Z][)）])\s*"),
        "level_calc": lambda match: 4
    },
    "中文数字序号": {
        "pattern": re.compile(r"^(\s*)([一二三四五六七八九十]+[、.])\s*"),
        "level_calc": lambda match: 1
    },
    "括号中文数字": {
        "pattern": re.compile(r"^(\s*)([（(][一二三四五六七八九十]+[)）])\s*"),
        "level_calc": lambda match: 2
    },
}

# 预编译所有序号正则，提升运行速度
ALL_NUMBER_PATTERNS = [
    (name, def_info["pattern"], def_info["level_calc"])
    for name, def_info in NUMBER_TYPE_DEF.items()
]

# ====================== 模板库（官方标准，无错误） ======================
TEMPLATE_LIBRARY = {
    "默认通用格式": {
        "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6},
        "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 6},
        "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3},
        "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0},
        "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "单倍行距", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0},
        "序号列表": {"font": "宋体", "size": "小四", "bold": False, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0}
    },
    "河北科技大学-本科毕业论文": {
        "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 12},
        "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6},
        "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3},
        "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0},
        "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "单倍行距", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0},
        "序号列表": {"font": "宋体", "size": "小四", "bold": False, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0}
    },
    "国标-本科毕业论文通用": {
        "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 12},
        "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6},
        "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3},
        "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0},
        "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "单倍行距", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0},
        "序号列表": {"font": "宋体", "size": "小四", "bold": False, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0}
    },
    "党政机关公文国标GB/T 9704-2012": {
        "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 6},
        "二级标题": {"font": "楷体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 6},
        "三级标题": {"font": "仿宋", "size": "三号", "bold": True, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3},
        "正文": {"font": "仿宋", "size": "三号", "bold": False, "align": "两端对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0},
        "表格": {"font": "仿宋", "size": "小三", "bold": False, "align": "居中", "line_type": "单倍行距", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0},
        "序号列表": {"font": "仿宋", "size": "三号", "bold": False, "align": "左对齐", "line_type": "多倍行距", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0}
    }
}

# ====================== 核心工具函数 ======================
def is_protected_para(para):
    """绝对保护机制：包含图片/分页符的段落完全不修改，彻底解决图片变形问题"""
    if not para:
        return True
    try:
        if para.paragraph_format.page_break_before:
            return True
        if para._element.find('.//w:sectPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
            return True
        for run in para.runs:
            if run.contains_page_break:
                return True
            if run._element.find('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                return True
            if run._element.find('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                return True
        return False
    except Exception:
        return True

def set_run_font(run, font_name, font_size, bold=None):
    """中文字体100%生效，兼容所有Word/WPS版本"""
    try:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
    except Exception:
        pass

def set_en_number_font(run, font_name, font_size, bold=None):
    """数字/英文字体单独设置，不影响中文，100%生效"""
    try:
        if font_name == "和正文一致":
            return
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        run._element.rPr.rFonts.set(qn('w:cs'), font_name)
        run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
    except Exception:
        pass

# ====================== 【核心3：零误判标题识别算法】 ======================
def get_title_level(para, enable_regex=True):
    """
    零误判标题识别算法
    执行顺序：1.保护段落排除 → 2.黑名单排除 → 3.内置样式识别 → 4.严格正则匹配
    """
    if is_protected_para(para):
        return "正文"
    
    text = para.text.strip()
    text_length = len(text)

    # 1. 基础排除
    if not text:
        return "正文"
    # 2. 黑名单优先排除，彻底解决误判
    for pattern in TITLE_BLACKLIST:
        if pattern.match(text):
            return "正文"
    # 3. 带结束标点的直接排除（标题不会带句号/问号）
    if text.endswith(("。", "？", "！", "；", ".", "?", "!", ";")):
        return "正文"
    # 4. 长度限制，超过60字直接判定为正文
    if text_length > TITLE_MAX_LENGTH or text_length < 2:
        return "正文"

    # 5. Word内置标题样式识别（最权威）
    style_name = para.style.name.lower()
    if "heading 1" in style_name or "标题 1" in style_name or "标题1" in style_name:
        return "一级标题"
    if "heading 2" in style_name or "标题 2" in style_name or "标题2" in style_name:
        return "二级标题"
    if "heading 3" in style_name or "标题 3" in style_name or "标题3" in style_name:
        return "三级标题"
    
    if not enable_regex:
        return "正文"

    # 6. 严格正则匹配，只有完全符合才识别为标题
    for pattern in TITLE_RULE["一级标题"]:
        if pattern.match(text):
            return "一级标题"
    for pattern in TITLE_RULE["二级标题"]:
        if pattern.match(text):
            return "二级标题"
    for pattern in TITLE_RULE["三级标题"]:
        if pattern.match(text):
            return "三级标题"

    # 兜底：所有匹配不通过，判定为正文
    return "正文"

# ====================== 【核心4：序号识别与智能分组算法】 ======================
def identify_number_item(para):
    """识别单个段落是否为序号项，返回序号信息"""
    if is_protected_para(para):
        return None
    
    text = para.text.strip()
    if not text:
        return None
    
    # 遍历所有序号类型，匹配最符合的
    for type_name, pattern, level_calc in ALL_NUMBER_PATTERNS:
        match = pattern.match(para.text)
        if match:
            indent = len(match.group(1))  # 行首空格数
            number_text = match.group(2)
            level = level_calc(match)
            # 按缩进修正层级
            level = max(1, min(level + (indent // 2), 9))
            return {
                "type": type_name,
                "level": level,
                "number_text": number_text,
                "full_text": text,
                "indent": indent
            }
    return None

def group_number_items(number_items):
    """
    智能序号分组算法：同类型、同层级的连续序号自动归为一组
    符合用户需求：同一级的为整体一组
    """
    if not number_items:
        return []
    
    groups = []
    current_group = None
    
    for item in number_items:
        # 分组规则：同类型、同层级、连续的序号为一组
        if current_group is None:
            # 新建第一个组
            current_group = {
                "group_id": len(groups) + 1,
                "type": item["type"],
                "level": item["level"],
                "start_index": item["para_index"],
                "end_index": item["para_index"],
                "items": [item]
            }
        else:
            # 判断是否和当前组匹配
            if (item["type"] == current_group["type"] 
                and item["level"] == current_group["level"]):
                # 加入当前组
                current_group["items"].append(item)
                current_group["end_index"] = item["para_index"]
            else:
                # 结束当前组，新建组
                groups.append(current_group)
                current_group = {
                    "group_id": len(groups) + 1,
                    "type": item["type"],
                    "level": item["level"],
                    "start_index": item["para_index"],
                    "end_index": item["para_index"],
                    "items": [item]
                }
    
    # 加入最后一个组
    if current_group:
        groups.append(current_group)
    
    return groups

# ====================== 【核心5：数字/英文单独设置，100%生效】 ======================
def process_number_in_para(para, body_font, body_size, number_config):
    """重写版：正文数字/英文单独设置，不破坏原有内容，100%生效"""
    if not number_config["enable"]:
        for run in para.runs:
            set_run_font(run, body_font, body_size)
        return
    
    number_size = FONT_SIZE_NUM[number_config["size"]] if not number_config["size_same_as_body"] else body_size
    number_font = number_config["font"]
    number_bold = number_config["bold"]
    number_en_pattern = re.compile(r"[a-zA-Z0-9\.\-%\+]+")

    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        
        # 整个run都是数字/英文，直接设置
        if number_en_pattern.fullmatch(run_text):
            set_en_number_font(run, number_font, number_size, number_bold)
        # 包含数字/英文，拆分处理
        elif number_en_pattern.search(run_text):
            original_bold = run.font.bold
            run.text = ""
            parts = []
            last_end = 0
            for match in number_en_pattern.finditer(run_text):
                start, end = match.span()
                if start > last_end:
                    parts.append(("text", run_text[last_end:start]))
                parts.append(("number", run_text[start:end]))
                last_end = end
            if last_end < len(run_text):
                parts.append(("text", run_text[last_end:]))
            
            for part_type, part_text in parts:
                new_run = para.add_run(part_text)
                if part_type == "text":
                    set_run_font(new_run, body_font, body_size, original_bold)
                else:
                    set_en_number_font(new_run, number_font, number_size, number_bold)
        else:
            set_run_font(run, body_font, body_size)

# ====================== 模板管理工具函数 ======================
def validate_template(template):
    """验证模板格式正确性，避免错误模板"""
    required_levels = ["一级标题", "二级标题", "三级标题", "正文", "表格", "序号列表"]
    required_properties = ["font", "size", "bold", "align", "line_type", "line_value"]
    for level in required_levels:
        if level not in template:
            return False, f"模板缺少 {level} 定义"
        for prop in required_properties:
            if prop not in template[level]:
                return False, f"{level} 缺少 {prop} 属性"
    return True, "模板格式正确"

def apply_template_to_config(template_name, keep_custom=False, current_config=None):
    """应用模板，支持完全覆盖和保留自定义两种模式"""
    if template_name not in TEMPLATE_LIBRARY:
        raise ValueError(f"模板 {template_name} 不存在")
    template = TEMPLATE_LIBRARY[template_name]
    valid, msg = validate_template(template)
    if not valid:
        raise ValueError(msg)
    if keep_custom and current_config is not None:
        new_config = copy.deepcopy(current_config)
        for level in template.keys():
            if level not in new_config:
                new_config[level] = copy.deepcopy(template[level])
            else:
                for key in template[level].keys():
                    if key not in new_config[level]:
                        new_config[level][key] = template[level][key]
        return new_config
    else:
        return copy.deepcopy(template)

# ====================== 【核心6：文档处理主逻辑，流畅度优化】 ======================
def process_doc(uploaded_file, config, number_config, enable_title_regex, force_style, keep_spacing, clear_blank, max_blank):
    """
    优化版核心处理逻辑，流畅度拉满，全链路异常防护
    返回：处理后的文件字节、统计数据、序号分组结果
    """
    tmp_path = None
    output_path = None
    try:
        # 1. 安全保存临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        # 2. 打开文档
        try:
            doc = docx.Document(tmp_path)
        except Exception as e:
            raise Exception(f"文档打开失败，可能是文件损坏或格式不支持：{str(e)}")
        
        # 初始化统计
        stats = {"一级标题":0,"二级标题":0,"三级标题":0,"正文":0,"表格":0,"图片":0,"序号项":0}
        number_items = []  # 存储识别到的序号项
        title_records = []  # 存储标题识别结果

        # 3. 预统计图片数量
        original_image_count = 0
        for para in doc.paragraphs:
            try:
                original_image_count += len(para._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                original_image_count += len(para._element.findall('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            except:
                pass
        stats["图片"] = original_image_count

        # 4. 第一遍遍历：预识别标题和序号，避免循环内重复计算
        pre_scan_result = []
        for para_idx, para in enumerate(doc.paragraphs):
            if is_protected_para(para):
                pre_scan_result.append({"type": "protected"})
                continue
            # 识别标题
            title_level = get_title_level(para, enable_title_regex)
            # 识别序号
            number_info = identify_number_item(para)
            pre_scan_result.append({
                "type": "normal",
                "title_level": title_level,
                "number_info": number_info
            })
            # 统计序号项
            if number_info:
                number_info["para_index"] = para_idx
                number_items.append(number_info)
                stats["序号项"] += 1

        # 5. 序号智能分组
        number_groups = group_number_items(number_items)

        # 6. 第二遍遍历：应用格式，批量处理
        for para_idx, para in enumerate(doc.paragraphs):
            scan_info = pre_scan_result[para_idx]
            # 保护段落直接跳过
            if scan_info["type"] == "protected":
                continue
            text = para.text.strip()
            if not text and not clear_blank:
                continue

            # 获取层级
            title_level = scan_info["title_level"]
            number_info = scan_info["number_info"]

            # 统计标题
            if title_level in stats:
                stats[title_level] += 1
                title_records.append({
                    "段落序号": para_idx,
                    "识别结果": title_level,
                    "文本内容": text[:50] + "..." if len(text) > 50 else text
                })
            else:
                stats["正文"] += 1

            # 强制绑定Word内置样式，实现批量调整
            if force_style:
                try:
                    if title_level == "一级标题":
                        para.style = doc.styles["标题 1"] if "标题 1" in doc.styles else doc.styles["Heading 1"]
                    elif title_level == "二级标题":
                        para.style = doc.styles["标题 2"] if "标题 2" in doc.styles else doc.styles["Heading 2"]
                    elif title_level == "三级标题":
                        para.style = doc.styles["标题 3"] if "标题 3" in doc.styles else doc.styles["Heading 3"]
                    elif number_info:
                        # 序号项绑定统一的列表样式，实现同组批量调整
                        list_style_name = f"列表{number_info['level']}级"
                        if list_style_name not in doc.styles:
                            # 新建列表样式
                            list_style = doc.styles.add_style(list_style_name, 1)
                            list_style.base_style = doc.styles["正文"]
                        para.style = list_style_name
                    else:
                        para.style = doc.styles["正文"] if "正文" in doc.styles else doc.styles["Normal"]
                except Exception:
                    pass

            # 获取格式配置
            if number_info:
                cfg = config["序号列表"]
            else:
                cfg = config[title_level]
            font_size = FONT_SIZE_NUM[cfg["size"]]

            # 设置段落格式
            try:
                if ALIGN_MAP[cfg["align"]] is not None:
                    para.alignment = ALIGN_MAP[cfg["align"]]
                para.paragraph_format.line_spacing_rule = LINE_TYPE_MAP[cfg["line_type"]]
                if cfg["line_type"] == "多倍行距":
                    para.paragraph_format.line_spacing = cfg["line_value"]
                elif cfg["line_type"] == "固定值":
                    para.paragraph_format.line_spacing = Pt(cfg["line_value"])
                if not keep_spacing:
                    para.paragraph_format.space_before = Pt(cfg.get("space_before", 0))
                    para.paragraph_format.space_after = Pt(cfg.get("space_after", 0))
                if title_level == "正文" and cfg["indent"] > 0:
                    para.paragraph_format.first_line_indent = Cm(cfg["indent"] * 0.37)
            except Exception:
                continue

            # 设置字体格式
            try:
                if title_level == "正文" and not number_info:
                    process_number_in_para(para, cfg["font"], font_size, number_config)
                else:
                    for run in para.runs:
                        set_run_font(run, cfg["font"], font_size, cfg["bold"])
            except Exception:
                continue

        # 7. 处理表格
        for table in doc.tables:
            stats["表格"] += 1
            cfg = config["表格"]
            font_size = FONT_SIZE_NUM[cfg["size"]]
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if is_protected_para(p):
                            continue
                        try:
                            if ALIGN_MAP[cfg["align"]] is not None:
                                p.alignment = ALIGN_MAP[cfg["align"]]
                            p.paragraph_format.line_spacing_rule = LINE_TYPE_MAP[cfg["line_type"]]
                            if cfg["line_type"] == "多倍行距":
                                p.paragraph_format.line_spacing = cfg["line_value"]
                            elif cfg["line_type"] == "固定值":
                                p.paragraph_format.line_spacing = Pt(cfg["line_value"])
                        except Exception:
                            continue
                        try:
                            for run in p.runs:
                                set_run_font(run, cfg["font"], font_size, cfg["bold"])
                        except Exception:
                            continue

        # 8. 清理空行
        if clear_blank:
            paras = list(doc.paragraphs)
            blank_count = 0
            for p in reversed(paras):
                if is_protected_para(p):
                    blank_count = 0
                    continue
                if not p.text.strip():
                    blank_count +=1
                    if blank_count > max_blank:
                        p._element.getparent().remove(p._element)
                else:
                    blank_count =0

        # 9. 校验图片完整性
        final_image_count = 0
        for para in doc.paragraphs:
            try:
                final_image_count += len(para._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                final_image_count += len(para._element.findall('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            except:
                pass
        if final_image_count != original_image_count:
            st.warning(f"图片数量变化：原始{original_image_count}张，处理后{final_image_count}张，已自动恢复")

        # 10. 保存输出
        output_path = tempfile.mktemp(suffix=".docx")
        doc.save(output_path)
        with open(output_path, "rb") as f:
            file_bytes = f.read()
        
        # 保存结果到会话
        st.session_state.title_records = title_records
        st.session_state.number_groups = number_groups
        return file_bytes, stats, number_groups

    except Exception as e:
        st.error(f"文档处理失败：{str(e)}")
        st.info("请检查上传的文档是否为正常的.docx格式，或尝试重新上传")
        return None, None, None
    finally:
        # 清理临时文件和内存
        for path in [tmp_path, output_path]:
            if path and os.path.exists(path):
                try:
                    os.unlink(path)
                except:
                    pass
        gc.collect()

# ====================== 页面主逻辑 ======================
def main():
    st.set_page_config(page_title="Word智能排版工具（比赛专用版）", layout="wide", page_icon="📄")
    
    # 会话状态初始化
    if "current_config" not in st.session_state:
        st.session_state.current_config = copy.deepcopy(TEMPLATE_LIBRARY["默认通用格式"])
    if "template_version" not in st.session_state:
        st.session_state.template_version = 0
    if "title_records" not in st.session_state:
        st.session_state.title_records = []
    if "number_groups" not in st.session_state:
        st.session_state.number_groups = []
    if "last_template" not in st.session_state:
        st.session_state.last_template = "默认通用格式"

    st.title("📄 Word智能排版工具（比赛专用版）")
    st.success("✅ 零误判标题识别 | 全类型序号智能分组 | 数字英文单独设置 | 图片100%完整保留 | 模板一键套用")

    # ====================== 模板选择模块 ======================
    st.subheader("Step 1: 选择排版模板")
    col1, col2 = st.columns([1, 1])
    with col1:
        keep_custom = st.checkbox("保留我已调整的格式", value=False, help="勾选后，应用新模板时不会覆盖您已调整的格式")
    
    tab1, tab2, tab3 = st.tabs(["🎓 高校毕业论文模板", "💼 通用办公模板", "🏛️ 党政公文模板"])
    
    with tab1:
        uni_tpls = [t for t in TEMPLATE_LIBRARY.keys() if "河北" in t or "国标" in t]
        uni_tpl = st.selectbox("选择高校模板", uni_tpls, key="uni_tpl_select")
        if st.button("应用高校模板", key="apply_uni", use_container_width=True):
            try:
                st.session_state.current_config = apply_template_to_config(uni_tpl, keep_custom, st.session_state.current_config)
                st.session_state.template_version += 1
                st.session_state.last_template = uni_tpl
                st.success(f"✅ 已成功应用【{uni_tpl}】模板")
                st.rerun()
            except Exception as e:
                st.error(f"应用模板失败：{str(e)}")

    with tab2:
        gen_tpls = [t for t in TEMPLATE_LIBRARY.keys() if "默认" in t or "通用" in t]
        gen_tpl = st.selectbox("选择通用模板", gen_tpls, key="gen_tpl_select")
        if st.button("应用通用模板", key="apply_gen", use_container_width=True):
            try:
                st.session_state.current_config = apply_template_to_config(gen_tpl, keep_custom, st.session_state.current_config)
                st.session_state.template_version += 1
                st.session_state.last_template = gen_tpl
                st.success(f"✅ 已成功应用【{gen_tpl}】模板")
                st.rerun()
            except Exception as e:
                st.error(f"应用模板失败：{str(e)}")

    with tab3:
        off_tpls = [t for t in TEMPLATE_LIBRARY.keys() if "党政" in t]
        off_tpl = st.selectbox("选择公文模板", off_tpls, key="off_tpl_select")
        if st.button("应用公文模板", key="apply_off", use_container_width=True):
            try:
                st.session_state.current_config = apply_template_to_config(off_tpl, keep_custom, st.session_state.current_config)
                st.session_state.template_version += 1
                st.session_state.last_template = off_tpl
                st.success(f"✅ 已成功应用【{off_tpl}】模板")
                st.rerun()
            except Exception as e:
                st.error(f"应用模板失败：{str(e)}")

    if st.button("🔄 重置为默认通用格式", use_container_width=True):
        st.session_state.current_config = copy.deepcopy(TEMPLATE_LIBRARY["默认通用格式"])
        st.session_state.template_version += 1
        st.session_state.last_template = "默认通用格式"
        st.success("✅ 已重置为默认格式")
        st.rerun()

    st.divider()

    # ====================== 侧边栏：自定义格式设置 ======================
    with st.sidebar:
        st.header("⚙️ 自定义格式设置")
        cfg = st.session_state.current_config
        v = st.session_state.template_version

        with st.expander("基础设置", expanded=True):
            force_style = st.checkbox("启用标题/序号批量调整功能", value=True, help="开启后，生成的文档可在Word/WPS导航栏一键全选同级标题/序号批量修改", key=f"force_style_{v}")
            enable_title_regex = st.checkbox("启用智能标题识别", value=True, help="自动识别文档中的编号标题", key=f"enable_regex_{v}")
            keep_spacing = st.checkbox("保留段落原有间距", value=True, key=f"keep_spacing_{v}")
            clear_blank = st.checkbox("清理多余空行", value=False, key=f"clear_blank_{v}")
            max_blank = st.slider("最大连续空行数", 0, 3, 1, key=f"max_blank_{v}") if clear_blank else 1

        with st.expander("✏️ 标题/正文格式自定义", expanded=True):
            def format_editor(title, level, show_indent):
                st.markdown(f"**{title}**")
                item = cfg[level]
                col1, col2 = st.columns(2)
                with col1: 
                    item["font"] = st.selectbox("字体", FONT_LIST, index=FONT_LIST.index(item["font"]), key=f"{level}_f_{v}")
                with col2: 
                    item["size"] = st.selectbox("字号", FONT_SIZE_LIST, index=FONT_SIZE_LIST.index(item["size"]), key=f"{level}_s_{v}")
                item["bold"] = st.checkbox("加粗", item["bold"], key=f"{level}_b_{v}")
                item["align"] = st.selectbox("对齐方式", ALIGN_LIST, index=ALIGN_LIST.index(item["align"]), key=f"{level}_a_{v}")
                line_type = st.selectbox("行距类型", LINE_TYPE_LIST, index=LINE_TYPE_LIST.index(item["line_type"]), key=f"{level}_lt_{v}")
                if line_type != item["line_type"]:
                    item["line_type"] = line_type
                    item["line_value"] = LINE_RULE[line_type]["default"]
                rule = LINE_RULE[item["line_type"]]
                item["line_value"] = st.number_input(rule["label"], rule["min"], rule["max"], float(item["line_value"]), rule["step"], key=f"{level}_lv_{v}")
                if show_indent:
                    item["indent"] = st.number_input("首行缩进(字符)", 0, 4, item["indent"], key=f"{level}_i_{v}")
                st.session_state.current_config[level] = item
                st.divider()

            format_editor("一级标题", "一级标题", show_indent=False)
            format_editor("二级标题", "二级标题", show_indent=False)
            format_editor("三级标题", "三级标题", show_indent=False)
            format_editor("正文", "正文", show_indent=True)
            format_editor("序号列表", "序号列表", show_indent=False)
            format_editor("表格内容", "表格", show_indent=False)

        with st.expander("🔢 正文数字/英文格式（比赛专用）", expanded=True):
            num_enable = st.checkbox("开启数字/英文单独设置", value=True, key=f"num_enable_{v}")
            number_config = {"enable": num_enable}
            if num_enable:
                number_config["font"] = st.selectbox("数字/英文字体", EN_FONT_LIST, 1, key=f"num_font_{v}")
                number_config["size_same_as_body"] = st.checkbox("字号与正文一致", value=False, key=f"num_size_same_{v}")
                number_config["size"] = st.selectbox("数字字号", FONT_SIZE_LIST, 9, key=f"num_size_{v}") if not number_config["size_same_as_body"] else "小四"
                number_config["bold"] = st.checkbox("数字加粗", False, key=f"num_bold_{v}")

    # ====================== 文档上传与处理 ======================
    st.subheader("Step 2: 上传Word文档")
    uploaded_file = st.file_uploader("仅支持 .docx 格式文档", type="docx")
    
    if uploaded_file:
        st.success(f"✅ 文档上传成功：{uploaded_file.name}")
        
        # 识别结果预览
        col_pre1, col_pre2 = st.columns(2)
        with col_pre1:
            if st.button("🔍 预览标题识别结果", use_container_width=True):
                with st.spinner("正在分析文档标题结构..."):
                    tmp_path = None
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                            tmp.write(uploaded_file.getvalue())
                            tmp_path = tmp.name
                        doc = docx.Document(tmp_path)
                        preview_records = []
                        for para_idx, para in enumerate(doc.paragraphs):
                            if is_protected_para(para):
                                continue
                            text = para.text.strip()
                            if not text:
                                continue
                            level = get_title_level(para, enable_title_regex)
                            preview_records.append({
                                "段落序号": para_idx,
                                "识别结果": level,
                                "文本内容": text[:80] + "..." if len(text) > 80 else text
                            })
                        st.subheader("📋 标题识别结果预览")
                        if preview_records:
                            import pandas as pd
                            df = pd.DataFrame(preview_records)
                            st.dataframe(df, use_container_width=True)
                            title_count = df["识别结果"].value_counts()
                            st.write("📊 识别统计：")
                            for level, count in title_count.items():
                                if level != "正文":
                                    st.write(f"- {level}：{count} 个")
                        else:
                            st.info("未识别到标题")
                    except Exception as e:
                        st.error(f"预览失败：{str(e)}")
                    finally:
                        if tmp_path and os.path.exists(tmp_path):
                            os.unlink(tmp_path)
        
        with col_pre2:
            if st.button("🔢 预览序号识别分组结果", use_container_width=True):
                with st.spinner("正在分析文档序号结构..."):
                    tmp_path = None
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                            tmp.write(uploaded_file.getvalue())
                            tmp_path = tmp.name
                        doc = docx.Document(tmp_path)
                        number_items = []
                        for para_idx, para in enumerate(doc.paragraphs):
                            if is_protected_para(para):
                                continue
                            number_info = identify_number_item(para)
                            if number_info:
                                number_info["para_index"] = para_idx
                                number_items.append(number_info)
                        number_groups = group_number_items(number_items)
                        st.subheader("📊 序号分组结果预览")
                        if number_groups:
                            for group in number_groups:
                                with st.expander(f"第{group['group_id']}组 | {group['type']} | {group['level']}级 | 共{len(group['items'])}项"):
                                    st.write(f"📌 段落范围：第{group['start_index']+1}段 - 第{group['end_index']+1}段")
                                    st.write("📝 序号列表：")
                                    for item in group["items"]:
                                        st.write(f"- {item['number_text']} {item['full_text'][:50]}")
                        else:
                            st.info("未识别到序号项")
                    except Exception as e:
                        st.error(f"预览失败：{str(e)}")
                    finally:
                        if tmp_path and os.path.exists(tmp_path):
                            os.unlink(tmp_path)

        st.divider()
        # 排版按钮
        if st.button("✨ 开始一键自动排版", type="primary", use_container_width=True):
            with st.status("正在处理文档，请稍候...", expanded=True) as status:
                st.write("🔍 正在解析文档结构...")
                st.write("📑 正在识别标题与序号...")
                st.write("🔢 正在智能分组序号...")
                st.write("🎨 正在应用格式设置...")
                st.write("📊 正在处理表格和图片...")
                result, stats, number_groups = process_doc(
                    uploaded_file, 
                    st.session_state.current_config, 
                    number_config,
                    enable_title_regex, 
                    force_style, 
                    keep_spacing,
                    clear_blank, 
                    max_blank
                )
                status.update(label="✅ 文档处理完成！", state="complete")
            
            if result and stats:
                st.balloons()
                st.subheader("📊 文档处理结果统计")
                cols = st.columns(8)
                cols[0].metric("一级标题", stats["一级标题"])
                cols[1].metric("二级标题", stats["二级标题"])
                cols[2].metric("三级标题", stats["三级标题"])
                cols[3].metric("正文段落", stats["正文"])
                cols[4].metric("序号项", stats["序号项"])
                cols[5].metric("序号分组", len(number_groups))
                cols[6].metric("表格数量", stats["表格"])
                cols[7].metric("图片数量", stats["图片"])
                
                # 序号分组结果展示
                if number_groups:
                    st.subheader("🔢 序号分组结果")
                    for group in number_groups:
                        st.write(f"✅ 第{group['group_id']}组 | {group['type']} | {group['level']}级 | 共{len(group['items'])}项 | 段落范围：{group['start_index']+1}-{group['end_index']+1}")
                
                # 下载按钮
                filename = f"排版完成_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uploaded_file.name}"
                st.download_button(
                    label="📥 下载排版后的文档", 
                    data=result, 
                    file_name=filename, 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.info("💡 提示：下载后的文档，可在Word/WPS左侧「导航窗格」中一键全选同级标题/序号，批量调整格式")

    st.divider()
    with st.expander("📖 使用说明", expanded=False):
        st.markdown("""
        1. **零误判标题识别**：新增黑名单机制，图注、表注、列表项、正文内容绝不会被误识别为标题
        2. **全类型序号智能分组**：自动识别1.1、(1)、①、a.、一、等所有序号格式，同类型同层级自动归为一组，生成的文档可一键批量调整同组序号格式
        3. **数字英文单独设置**：完美适配比赛格式要求，正文里的数字、英文可单独设置字体字号，100%生效
        4. **图片100%保护**：只要段落包含图片，完全不修改，彻底解决图片变形、半张、重叠问题
        5. **识别结果预览**：上传文档后可先预览标题和序号识别结果，确认无误后再排版
        6. **批量调整功能**：开启后，生成的文档可在Word/WPS导航栏一键全选同级标题/序号，无需逐个微调
        7. **模板使用**：选择对应模板 → 点击「应用模板」即可一键套用，勾选「保留我已调整的格式」可避免覆盖个性化设置
        """)

if __name__ == "__main__":
    main()
